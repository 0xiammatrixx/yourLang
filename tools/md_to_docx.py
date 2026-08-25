"""Convert docs/paper.md into a Word document for submission.

The paper is maintained in Markdown (docs/paper.md) and exported to .docx
for submission. Handles: title/headings, paragraphs, bullet lists, numbered
lists (literal numbers preserved), pipe tables, and fenced code blocks
(including the Mermaid diagram).

Usage:
    .venv/bin/python tools/md_to_docx.py [INPUT.md] [OUTPUT.docx]
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

INLINE = re.compile(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)")


def add_inline_runs(paragraph, text):
    """Add runs to *paragraph* honouring **bold**, *italic*, and `code`."""
    for token in INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def cell_text(cell):
    """Strip markdown markers from a table cell; return (text, all_bold)."""
    text = cell.strip()
    all_bold = False
    if text.startswith("**") and text.endswith("**") and len(text) > 4:
        text = text[2:-2]
        all_bold = True
    return re.sub(r"`|\*", "", text), all_bold


def add_table(doc, rows):
    """rows: list of lists of raw markdown cell strings (header included)."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            text, all_bold = cell_text(row[j]) if j < len(row) else ("", False)
            para = cells[j].paragraphs[0]
            run = para.add_run(text)
            run.bold = all_bold or i == 0
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    def add_para(text, style=None):
        p = doc.add_paragraph(style=style)
        add_inline_runs(p, text)
        return p

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line separates blocks.
        if not stripped:
            i += 1
            continue

        # Fenced code block (includes the Mermaid diagram).
        if stripped.startswith("```"):
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph()
                run = p.add_run(lines[i])
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                i += 1
            i += 1  # closing fence
            continue

        # Headings.
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
            i += 1
            continue

        # Pipe table: collect consecutive table lines.
        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # Skip the |---|---| separator row.
                if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells):
                    rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        # Bullet list items (with hanging continuation lines merged).
        bullet = re.match(r"^([-*]) (.*)$", stripped)
        if bullet:
            indent = len(line) - len(line.lstrip())
            text = bullet.group(2)
            i += 1
            while i < n and lines[i].startswith("  ") and lines[i].strip():
                if re.match(r"^\s*([-*] |\d+\.)", lines[i]):
                    break  # nested list item: handle in its own iteration
                text += " " + lines[i].strip()
                i += 1
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            add_para(text, style=style)
            continue

        # Numbered list items: keep the literal number (Word auto-numbering
        # would restart unpredictably across sections).
        numbered = re.match(r"^\d+\. (.*)$", stripped)
        if numbered:
            text = numbered.group(1)
            i += 1
            while i < n and lines[i].startswith("  ") and lines[i].strip():
                if re.match(r"^\s*([-*] |\d+\.)", lines[i]):
                    break
                text += " " + lines[i].strip()
                i += 1
            add_para(f"{numbered.group(0).split('.', 1)[0]}. {text}")
            continue

        # Plain paragraph.
        add_para(stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    md = Path(sys.argv[1] if len(sys.argv) > 1 else "docs/paper.md")
    out = Path(
        sys.argv[2]
        if len(sys.argv) > 2
        else "Semantic Boundary NL-Database Execution.docx"
    )
    convert(md, out)
    print(f"Wrote {out} ({out.stat().st_size} bytes) from {md}")


if __name__ == "__main__":
    main()
